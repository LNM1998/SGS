from flask import Blueprint, request, render_template, redirect, url_for, send_file, json
from flask_login import login_required
from models import db, Remito, ArticuloRemito
from datetime import date
from docx import Document
import os

remitos_bp = Blueprint('remitos', __name__)

UPLOAD_FOLDER = os.path.abspath('uploads/remitos')

@remitos_bp.route('/remitos')
@login_required
def remitos():
    remitos = Remito.query.all()
    return render_template('remitos.html', remitos=remitos)

@remitos_bp.route('/nuevo_remito', methods=['GET', 'POST'])
@login_required
def nuevo_remito():
    if request.method == 'POST':
        # Verifica que todos los campos requeridos estén presentes
        required_fields = ['remitente', 'destinatario', 'fecha']
        for field in required_fields:
            if field not in request.form or not request.form[field]:
                return f"Error: El campo '{field}' es requerido.", 400

        # Obtén los datos del formulario
        remitente = request.form['remitente']
        destinatario = request.form['destinatario']
        fecha = request.form['fecha']
        observacion = request.form.get('observacion', '')

        # Convierte la fecha a un objeto date
        try:
            fecha = date.fromisoformat(fecha) if fecha else None
        except ValueError:
            return "Error: Formato de fecha inválido. Use YYYY-MM-DD.", 400

        # Crea el remito
        remito = Remito(
            tipo_remito='Entrega',
            remitente=remitente,
            destinatario=destinatario,
            fecha=fecha,
            observacion=observacion,
            estado='pendiente'
        )
        db.session.add(remito)
        db.session.commit()  # Guarda el remito para obtener su ID

        # Procesar los artículos
        cantidades = request.form.getlist('cantidad[]')
        numeros_serie = request.form.getlist('numero_serie[]')
        descripciones = request.form.getlist('descripcion[]')
        observaciones_articulo = request.form.getlist('observacion[]')

        if not cantidades or not numeros_serie or not descripciones:
            return "Error: Debe cargar al menos un artículo.", 400

        for cantidad, numero_serie, descripcion, observacion in zip(cantidades, numeros_serie, descripciones, observaciones_articulo):
            try:
                cantidad = int(cantidad)
            except ValueError:
                return "Error: La cantidad debe ser un número entero.", 400

            # Crear el artículo
            articulo = ArticuloRemito(
                remito_id=remito.id,  # Asociar el artículo al remito
                cantidad=cantidad,
                numero_serie=numero_serie,
                descripcion=descripcion,
                observacion=observacion
            )
            db.session.add(articulo)

        db.session.commit()  # Guardar los artículos

        # Generar el documento DOCX
        try:
            archivo_path = llenar_plantilla(remito)
            remito.archivo_remito = archivo_path
            db.session.commit()
        except Exception as e:
            return f"Error al generar el archivo DOCX: {str(e)}", 500

        return """
        <script>
            window.opener.location.reload();  // Recargar la página principal
            window.close();  // Cerrar la ventana emergente
        </script>
        """
    return render_template('nuevo_remito.html')

@remitos_bp.route('/editar_remito/<int:id>', methods=['GET'])
@login_required
def editar_remito(id):
    remito = Remito.query.get_or_404(id)

    return render_template('editar_remito.html', remito=remito)

@remitos_bp.route('/actualizar_remito/<int:id>', methods=['POST'])
@login_required
def actualizar_estado(id):
    remito = Remito.query.get(id)

    if not remito:
        return "Remito no encontrado", 404

    if request.method == 'POST':
        # Actualizar campos del remito
        remito.tipo_remito = request.form.get('tipo_remito') or remito.tipo_remito
        remito.fecha_str = request.form.get('fecha')
        remito.fecha = date.fromisoformat(remito.fecha_str) if remito.fecha_str else remito.fecha
        remito.remitente = request.form.get('remitente') or remito.remitente
        remito.destinatario = request.form.get('destinatario') or remito.destinatario
        remito.observacion = request.form.get('observacion') or remito.observacion
        remito.estado = request.form.get('estado') or remito.estado

        # Procesar los artículos
        cantidades = request.form.getlist('cantidad[]')
        numeros_serie = request.form.getlist('numero_serie[]')
        descripciones = request.form.getlist('descripcion[]')
        observaciones_articulo = request.form.getlist('observacion_articulo[]')

        # Eliminar artículos antiguos
        for articulo in remito.articulos:
            db.session.delete(articulo)

        # Agregar nuevos artículos
        for cantidad, numero_serie, descripcion, observacion in zip(cantidades, numeros_serie, descripciones, observaciones_articulo):
            try:
                cantidad = int(cantidad)
            except ValueError:
                return "Error: La cantidad debe ser un número entero.", 400

            # Crear el artículo
            articulo = ArticuloRemito(
                remito_id=remito.id,
                cantidad=cantidad,
                numero_serie=numero_serie,
                descripcion=descripcion,
                observacion=observacion
            )
            db.session.add(articulo)

        # Acciones según el estado del remito
        if remito.estado.lower() == 'devuelto':
            remito.tipo_remito = 'Devolución'
            remito.archivo_remito = generar_remito_devolucion(remito)

        elif remito.estado.lower() == 'pendiente':
            remito.tipo_remito = 'Entrega'
            remito.archivo_remito = llenar_plantilla(remito)

        elif remito.estado.lower() == 'entrega definitiva':
            remito.tipo_remito = 'Entrega'
            remito.archivo_remito = llenar_plantilla(remito)

        db.session.commit()
        return """
            <script>
                window.opener.location.reload();  // Recargar la página principal
                window.close();  // Cerrar la ventana emergente
            </script>
            """
    
    # Redirigir a la lista de remitos después de actualizar
    return redirect(url_for('remitos', id=id))


@remitos_bp.route('/eliminar_remito/<int:id>', methods=['POST'])
@login_required
def eliminar_remito(id):
    remito = Remito.query.get(id)
    if remito:
        # Eliminar los artículos asociados al remito antes de eliminarlo
        ArticuloRemito.query.filter_by(remito_id=id).delete()

        db.session.delete(remito)
        db.session.commit()
    return redirect('/remitos')

@remitos_bp.route('/descargar/<int:id>')
@login_required
def descargar(id):
    # Si es ID 0, buscar el archivo en blanco en la carpeta 'uploads'
    if id == 0:
        archivo_path = os.path.join('uploads', 'remito_blanco.docx')
        
        if not os.path.exists(archivo_path):
            return "El remito en blanco no existe", 404
        
        return send_file(archivo_path, as_attachment=True)

    # Obtener el remito desde la base de datos
    remito = Remito.query.get_or_404(id)

    # Verificar si el remito ya tiene un archivo generado
    if not remito.archivo_remito or not os.path.exists(remito.archivo_remito):
        archivo_path = llenar_plantilla(remito)
        
        # Guardar la nueva ruta en la base de datos
        remito.archivo_remito = archivo_path
        db.session.commit()

    # Si el remito está marcado como "Devuelto", generar el remito de devolución
    if remito.estado.lower() == "devuelto":
        archivo_devolucion = generar_remito_devolucion(remito)
        
        # Guardar la ruta del remito de devolución en la base de datos
        remito.archivo_remito_devolucion = archivo_devolucion
        db.session.commit()

        return send_file(archivo_devolucion, as_attachment=True)

    return send_file(remito.archivo_remito, as_attachment=True)

def llenar_plantilla(remito):
    # Ruta de la plantilla
    template_dir = os.path.abspath(UPLOAD_FOLDER)
    template_name = "remito_entrega.docx"
    plantilla_path = os.path.join(template_dir, template_name)

    # Verificar si la plantilla existe
    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"La plantilla '{plantilla_path}' no existe.")

    # Crear directorio de salida si no existe
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

    # Cargar la plantilla
    doc = Document(plantilla_path)

    # Reemplazar datos generales del remito
    placeholders = {
        "{tipo_remito}": remito.tipo_remito or "Entrega",
        "{fecha}": remito.fecha.strftime("%d/%m/%Y") if remito.fecha else "Sin fecha",
        "{remitente}": remito.remitente or "Sin remitente",
        "{destinatario}": remito.destinatario or "Sin destinatario"
    }

    def reemplazar_texto(elemento):
        for placeholder, valor in placeholders.items():
            if placeholder in elemento.text:
                elemento.text = elemento.text.replace(placeholder, valor)

    for p in doc.paragraphs:
        reemplazar_texto(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_texto(cell)

    # Buscar la tabla donde se insertarán los artículos
    tabla_articulos = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{tabla_articulos}" in cell.text:
                    tabla_articulos = table
                    break
        if tabla_articulos:
            break

    if tabla_articulos:
        # Limpiar la celda donde estaba el marcador
        tabla_articulos.cell(1, 0).text = ""

        # Insertar los artículos en la tabla
        for articulo in remito.articulos:
            fila = tabla_articulos.add_row().cells
            fila[0].text = str(articulo.cantidad)
            fila[1].text = articulo.numero_serie
            fila[2].text = articulo.descripcion
            fila[3].text = articulo.observacion if articulo.observacion else "-"

    # Guardar el archivo generado
    archivo_path = os.path.join(UPLOAD_FOLDER, f"remito_entrega_{remito.destinatario}_{remito.id}.docx")
    doc.save(archivo_path)

    return archivo_path

def generar_remito_devolucion(remito):
    template_dir = os.path.abspath(UPLOAD_FOLDER)
    template_name = "remito_devolucion.docx"
    plantilla_path = os.path.join(template_dir, template_name)

    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"La plantilla de devolución '{plantilla_path}' no existe.")

    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

    doc = Document(plantilla_path)

    placeholders = {
        "{tipo_remito}": "Devolución",
        "{fecha}": remito.fecha.strftime("%d/%m/%Y") if remito.fecha else "Sin fecha",
        "{remitente}": remito.remitente or "Sin remitente",
        "{destinatario}": remito.destinatario or "Sin destinatario"
    }

    def reemplazar_texto(elemento):
        for placeholder, valor in placeholders.items():
            if placeholder in elemento.text:
                elemento.text = elemento.text.replace(placeholder, valor)

    for p in doc.paragraphs:
        reemplazar_texto(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_texto(cell)

    tabla_articulos = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{tabla_articulos}" in cell.text:
                    tabla_articulos = table
                    break
        if tabla_articulos:
            break

    if tabla_articulos:
        tabla_articulos.cell(1, 0).text = ""

        for articulo in remito.articulos:
            fila = tabla_articulos.add_row().cells
            fila[0].text = str(articulo.cantidad)
            fila[1].text = articulo.numero_serie
            fila[2].text = articulo.descripcion
            fila[3].text = articulo.observacion if articulo.observacion else "-"

    archivo_path = os.path.join(UPLOAD_FOLDER, f"remito_devolucion_{remito.destinatario}_{remito.id}.docx")
    doc.save(archivo_path)

    return archivo_path