from flask import Blueprint, request, render_template, redirect, current_app, send_file, send_from_directory, Response, jsonify
from flask_login import login_required, current_user
import unicodedata
from sqlalchemy import func
from models import db, Notebook
from .auth import lectura_allowed, basico_required
from .logs import logger
from docx import Document
import os
from datetime import date

notebooks_bp = Blueprint('notebooks', __name__)

UPLOAD_FOLDER = os.path.abspath('uploads/actas')

def quitar_acentos(cadena):
    return ''.join(
        c for c in unicodedata.normalize('NFD', cadena)
        if unicodedata.category(c) != 'Mn'
    )

def safe_log(action, entity, details, user=None):
    """Función de logging ultra-resistente a fallos"""
    try:
        user = user or (current_user.username if current_user.is_authenticated else 'system')
        details_str = str(details)

        current_app.logger.info(f"Intentando registrar log: {action}, {entity}, {details_str}, {user}")

        if hasattr(logger, 'is_initialized') and logger.is_initialized:
            logger.log(action, entity, details_str, user)
        else:
            current_app.logger.warning("Logger personalizado no inicializado")
    except Exception as e:
        current_app.logger.error(f"Fallo en safe_log: {str(e)}")

@notebooks_bp.route('/notebooks', methods=['GET'])
@login_required
@lectura_allowed
def notebooks():
    filtro = request.args.get('filtro', '')
    valor = request.args.get('valor', '')
    filtro_estado_notebook = request.args.get('estado', '')
    filtro_modelo = request.args.get('modelo', '')
    filtro_direccion = request.args.get('direccion', '')

    # Construimos la consulta base
    consulta = Notebook.query

    if filtro_estado_notebook:
        consulta = consulta.filter(Notebook.estado == filtro_estado_notebook)

    if filtro_modelo:
        consulta = consulta.filter(Notebook.modelo == filtro_modelo)

    if filtro_direccion:
        consulta = consulta.filter(Notebook.direccion == filtro_direccion)
        
    if filtro and valor:
        valor_normalizado = quitar_acentos(valor).lower()

        if filtro == 'inventario':
            consulta = consulta.filter(Notebook.inventario.contains(valor))
        elif filtro == 'numero_serie':
            consulta = consulta.filter(Notebook.numero_serie.contains(valor))
        elif filtro == 'descripcion':
            consulta = consulta.filter(Notebook.descripcion.contains(valor))
        elif filtro == 'usuario':
            valor_normalizado = quitar_acentos(valor).lower()
            consulta = consulta.filter(
                func.lower(func.replace(func.replace(func.replace(
                    func.replace(func.replace(func.replace(func.replace(
                        Notebook.usuario,
                        'á', 'a'),
                        'é', 'e'),
                        'í', 'i'),
                        'ó', 'o'),
                        'ú', 'u'),
                        'Á', 'a'),
                        'É', 'e')
                ).contains(valor_normalizado)
            )

    notebooks = consulta.all()
    return render_template('notebooks.html', notebooks=notebooks, filtro=filtro, valor=valor, filtro_estado_notebook=filtro_estado_notebook, filtro_modelo=filtro_modelo, filtro_direccion=filtro_direccion)

@notebooks_bp.route('/agregar_notebook', methods=['GET','POST'])
@login_required
@basico_required
def agregar_notebook():
    if request.method == 'POST':
        try:
            form_data = {
                'modelo' : request.form.get('modelo', '').strip(),
                'inventario' : request.form.get('inventerio', '').strip(),
                'numero_serie' : request.form.get('numero_serie', '').strip(),
                'estado' : request.form.get('estado', '').strip(),
                'usuario' : request.form.get('usuario', '').strip(),
                'direccion' : request.form.get('direccion', '').strip(),
                'descripcion' : request.form.get('descripcion', '').strip(),
                'fecha' : request.form.get('fecha', '').strip(),
            }

            fecha = date.fromisoformat(form_data['fecha']) if form_data['fecha'] else None
        
            notebook_existente = Notebook.query.filter_by(inventario=form_data['inventario']).first()
            if notebook_existente:
                log_data = {
                    'inventario': form_data['inventario'],
                    'intento': form_data,
                    'mensaje': 'Notebook ya existe en la base de datos'
                }
                safe_log('conflict', 'notebook', log_data, current_user.username)
                return """
                <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
                <script>
                    document.addEventListener("DOMContentLoaded", function() {
                        Swal.fire({
                            title: 'Error',
                            text: 'El número de inventario ya existe.',
                            icon: 'error',
                            confirmButtonText: 'Aceptar'
                        }).then(() => {
                            window.opener.location.reload();  // Recargar la página principal
                            window.close();
                        });
                    });
                </script>
                """

            nueva_notebook = Notebook(
                tipo_remito='Entrega',
                modelo=form_data['modelo'],
                inventario=form_data['inventario'], 
                numero_serie=form_data['numero_serie'], 
                estado=form_data['estado'], 
                usuario=form_data['usuario'], 
                fecha=fecha,
                direccion=form_data['direccion'],
                descripcion=form_data['descripcion'] or None
            )

            db.session.add(nueva_notebook)
            db.session.commit()

            archivo_path = llenar_plantilla_acta(nueva_notebook)
            nueva_notebook.archivo_notebook = archivo_path

            safe_log(
                action='create',
                entity='notebook',
                details={
                    'id': nueva_notebook.id,
                    'modelo' : nueva_notebook.modelo,
                    'inventario' : nueva_notebook.inventario,
                    'numero_serie' : nueva_notebook.numero_serie,
                    'estado' : nueva_notebook.estado,
                    'usuario' : nueva_notebook.usuario,
                    'fecha' : nueva_notebook.fecha,
                    'direccion' : nueva_notebook.direccion,
                    'descripcion' : nueva_notebook.descripcion
                },
                user=current_user.username
            )
            return """
            <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
            <script>
                document.addEventListener("DOMContentLoaded", function() {
                    Swal.fire({
                        title: 'Un Éxito',
                        text: 'Notebook agregada correctamente.',
                        icon: 'success',
                        confirmButtonText: 'Aceptar'
                    }).then(() => {
                        window.opener.location.reload();  // Recargar la página principal
                        window.close();  // Cerrar la ventana emergente
                    });
                });
            </script>
            """
        except Exception as e:
            db.session.rollback()
            
            current_app.logger.error(f"Error en agregar_equipo: {str(e)}")
            safe_log(
                action='error',
                entity='notebook',
                details={'error': str(e), 'operacion': 'agregar_notebook'},
                user=current_user.username
            )

            return """
            <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
            <script>
                document.addEventListener("DOMContentLoaded", function() {
                    Swal.fire({
                        title: 'Error',
                        text: 'Ocurrió un error al agregar el equipo.',
                        icon: 'error',
                        confirmButtonText: 'Aceptar'
                    }).then(() => {
                        window.opener.location.reload();
                        window.close();
                    });
                });
            </script>
            """
        
    return render_template('agregar_notebook.html')

@notebooks_bp.route('/eliminar_notebook/<int:id>', methods=['POST'])
@login_required
@basico_required
def eliminar_notebook(id):
    try:
        notebook = Notebook.query.get(id)

        safe_log(
            action='delete',
            entity='notebook',
            details={
                'id': id,
                'modelo' : notebook.modelo,
                'inventario' : notebook.inventario,
                'numero_serie' : notebook.numero_serie,
                'estado' : notebook.estado,
                'usuario' : notebook.usuario,
                'fecha' : notebook.fecha,
                'direccion' : notebook.direccion,
                'descripcion' : notebook.descripcion
            },
            user=current_user.username
        )

        db.session.delete(notebook)
        db.session.commit()

        return redirect('/notebooks')
    
    except Exception as e:
        db.session.rollback()
        safe_log(
            action='error',
            entity='notebook',
            details={
                'operation': 'delete',
                'error': str(e),
                'notebook_id': id
            },
            user=current_user.username
        )
        return "Error al eliminar la notebook", 500

@notebooks_bp.route('/editar_notebook/<int:id>', methods=['GET'])
@login_required
@basico_required
def editar_notebook(id):
    notebook = Notebook.query.get(id)
    return render_template('editar_notebook.html', notebook=notebook)

@notebooks_bp.route('/actualizar_notebook/<int:id>', methods=['POST'])
@login_required
@basico_required
def actualizar_notebook(id):
    notebook = Notebook.query.get(id)

    if request.method == 'POST':
        try:
            datos_antiguos = {
                'modelo' : notebook.modelo,
                'inventario' : notebook.inventario,
                'numero_serie' : notebook.numero_serie,
                'estado' : notebook.estado,
                'usuario' : notebook.usuario,
                'fecha' : notebook.fecha,
                'direccion' : notebook.direccion,
                'descripcion' : notebook.descripcion
            }

            notebook.modelo = request.form['modelo']
            notebook.inventario = request.form['inventario']
            notebook.numero_serie = request.form['numero_serie']
            notebook.estado = request.form['estado']
            notebook.usuario = request.form['usuario']
            notebook.direccion = request.form['direccion']
            notebook.descripcion = request.form['descripcion']

            notebook.fecha = request.form['fecha']

            if notebook.fecha:
                notebook.fecha = date.fromisoformat(notebook.fecha)
            else:
                notebook.fecha = None  # Si el campo está vacío, guarda NULL en la base de datos

            if notebook.estado.lower() == "fisica":
                if notebook.archivo_notebook and os.path.exists(notebook.archivo_notebook):
                    try:
                        os.remove(notebook.archivo_notebook)
                    except Exception as e:
                        print(f"No se pudo eliminar el archivo anterior: {e}")

                # 1. Generar el acta con los datos actuales
                archivo_devolucion = generar_acta_devolucion(notebook)

                # 2. Guardar la ruta del acta de devolución en la base de datos
                notebook.archivo_notebook = archivo_devolucion

                # 4. Vaciar los campos
                notebook.usuario = ""
                notebook.direccion = ""
                notebook.descripcion = ""

                # 5. Guardar cambios en la base de datos
                db.session.commit()

                safe_log(
                    action='update',
                    entity='notebook',
                    details={
                        'id': id,
                        'cambios': {
                            'antes': datos_antiguos,
                            'despues': {
                                'modelo' : notebook.modelo,
                                'inventario' : notebook.inventario,
                                'numero_serie' : notebook.numero_serie,
                                'estado' : notebook.estado,
                                'usuario' : notebook.usuario,
                                'fecha' : notebook.fecha,
                                'direccion' : notebook.direccion,
                                'descripcion' : notebook.descripcion
                            }
                        }
                    },
                    user=current_user.username
                )

                # 6. Retornar HTML con descarga automática y SweetAlert
                return f"""
                <html>
                <head>
                    <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
                </head>
                <body>
                    <script>
                        document.addEventListener("DOMContentLoaded", function() {{
                            Swal.fire({{
                                title: 'Un Éxito',
                                text: 'Notebook actualizada correctamente... Descargando acta de devolución...',
                                icon: 'success',
                                confirmButtonText: 'Aceptar'
                            }}).then(() => {{
                                // Iniciar descarga
                                const link = document.createElement('a');
                                link.href = '/ruta_de_descarga_directa/{notebook.id}';
                                link.download = '';
                                document.body.appendChild(link);
                                link.click();
                                document.body.removeChild(link);

                                // Recargar y cerrar ventana
                                setTimeout(() => {{
                                    window.opener.location.reload();
                                    window.close();
                                }}, 1000);
                            }});
                        }});
                    </script>
                </body>
                </html>
                """

            else:
                if notebook.archivo_notebook and os.path.exists(notebook.archivo_notebook):
                    try:
                        os.remove(notebook.archivo_notebook)
                    except Exception as e:
                        print(f"No se pudo eliminar el archivo anterior: {e}")

                archivo_entrega = llenar_plantilla_acta(notebook)
        
                # Guardar la nueva ruta en la base de datos
                notebook.archivo_notebook = archivo_entrega
                db.session.commit()

            safe_log(
                action='update',
                entity='notebook',
                details={
                    'id': id,
                    'cambios': {
                        'antes': datos_antiguos,
                        'despues': {
                            'modelo' : notebook.modelo,
                            'inventario' : notebook.inventario,
                            'numero_serie' : notebook.numero_serie,
                            'estado' : notebook.estado,
                            'usuario' : notebook.usuario,
                            'fecha' : notebook.fecha,
                            'direccion' : notebook.direccion,
                            'descripcion' : notebook.descripcion
                        }
                    }
                },
                user=current_user.username
            )
            return """
            <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
            <script>
                document.addEventListener("DOMContentLoaded", function() {
                    Swal.fire({
                        title: 'Un Éxito',
                        text: 'Notebook actualizada correctamente.',
                        icon: 'success',
                        confirmButtonText: 'Aceptar'
                    }).then(() => {
                        window.opener.location.reload();  // Recargar la página principal
                        window.close();  // Cerrar la ventana emergente
                    });
                });
            </script>
            """

        except Exception as e:
            db.session.rollback()
            safe_log(
                action='error',
                entity='notebook',
                details={
                    'operation': 'update',
                    'error': str(e),
                    'notebook_id': id
                },
                user=current_user.username
            )
            return "Error al actualizar la notebook", 500

    return redirect('/notebooks')

@notebooks_bp.route('/notebooks/ver_acta/<int:id>')
@login_required
@lectura_allowed
def ver_acta(id):
    notebook = Notebook.query.get_or_404(id)
    return render_template('acta.html', notebook=notebook)

@notebooks_bp.route('/ruta_de_descarga_directa/<int:id>')
@login_required
@lectura_allowed
def ruta_de_descarga_directa(id):
    notebook = Notebook.query.get_or_404(id)

    if notebook.archivo_notebook and os.path.exists(notebook.archivo_notebook):
        return send_file(notebook.archivo_notebook, as_attachment=True)
    else:
        return "Archivo no encontrado", 404

@notebooks_bp.route('/descargar_acta/<int:id>')
@login_required
@lectura_allowed
def descargar_acta(id):
    # Obtener el remito desde la base de datos
    notebook = Notebook.query.get_or_404(id)

    # Verificar si notebook ya tiene un archivo generado
    if not notebook.archivo_notebook or not os.path.exists(notebook.archivo_notebook):
        archivo_path = llenar_plantilla_acta(notebook)
        
        # Guardar la nueva ruta en la base de datos
        notebook.archivo_notebook = archivo_path
        db.session.commit()

    # Si el remito está marcado como "fisica", generar el remito de devolución
    if notebook.estado.lower() == "fisica":
        archivo_devolucion = generar_acta_devolucion(notebook)
        
        # Guardar la ruta del acta de devolución en la base de datos
        notebook.archivo_notebook = archivo_devolucion
        db.session.commit()

        return send_file(archivo_devolucion, as_attachment=True)

    return send_file(notebook.archivo_notebook, as_attachment=True)

def llenar_plantilla_acta(notebook):
    # Ruta de la plantilla
    template_dir = os.path.abspath(UPLOAD_FOLDER)
    template_name = "acta_entrega.docx"
    plantilla_path = os.path.join(template_dir, template_name)

    # Verificar si la plantilla existe
    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"La plantilla '{plantilla_path}' no existe.")

    # Crear directorio de salida si no existe
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

    # Cargar la plantilla
    doc = Document(plantilla_path)

    # Determinar el modelo formateado
    if notebook.modelo == 'dell i5':
        modelo_formateado = "Dell Latitude 3520 i5"
    elif notebook.modelo == 'dell i7':
        modelo_formateado = "Dell Latitude 5520 i7"
    elif notebook.modelo == 'dell i7 diseno':
        modelo_formateado = "Dell Precision 3561 i7"
    elif notebook.modelo == 'exo i5':
        modelo_formateado = "Exo Smart Pro Q6 i5"
    elif notebook.modelo == 'exo i7':
        modelo_formateado = "Exo Smart Pro Q6 i7"
    else:
        modelo_formateado = "Sin modelo"

    # Dirección formateada
    if notebook.direccion == 'intervencion':
        direccion_formateada = "Intervención"
    elif notebook.direccion == 'sistemas':
        direccion_formateada = "Dirección General de Sistemas Informáticos"
    elif notebook.direccion == 'fomento':
        direccion_formateada = "Dirección Nacional de Fomento y Desarrollo"
    elif notebook.direccion == 'administracion':
        direccion_formateada = "Dirección General de Administración"
    elif notebook.direccion == 'rrhh':
        direccion_formateada = "Dirección General de Recursos Humanos"
    elif notebook.direccion == 'juridicos':
        direccion_formateada = "Dirección General de Asuntos Jurídicos y Regulatorios"
    elif notebook.direccion == 'planificacion':
        direccion_formateada = "Dirección Nacional de Planificación y Convergencia"
    elif notebook.direccion == 'control':
        direccion_formateada = "Dirección Nacional de Control y Fiscalización"
    elif notebook.direccion == 'postales':
        direccion_formateada = "Dirección Nacional de Servicios Postales"
    elif notebook.direccion == 'institucionales':
        direccion_formateada = "Dirección General de Asuntos Institucionales"
    elif notebook.direccion == 'audiovisuales':
        direccion_formateada = "Dirección Nacional de Servicios Audiovisuales"
    elif notebook.direccion == 'competencia':
        direccion_formateada = "Dirección Nacional de Desarrollo de la Competencia en Redes"
    elif notebook.direccion == 'autorizaciones':
        direccion_formateada = "Dirección Nacional de Autorizaciones y Registros TIC"
    elif notebook.direccion == 'delegaciones':
        direccion_formateada = "Dirección Nacional de Atención de Usuarios"
    elif notebook.direccion == 'auditoria':
        direccion_formateada = "Unidad de Auditoría Interna"
    elif notebook.direccion == 'ccte':
        direccion_formateada = "Centro De Comprobación Técnica De Emisiones"
    else:
        direccion_formateada = "Sin dirección"


    # Diccionario de placeholders
    placeholders = {
        "{fecha}": notebook.fecha.strftime("%d/%m/%Y") if notebook.fecha else "Sin fecha",
        "{usuario}": notebook.usuario or "Sin usuario",
        "{direccion}": direccion_formateada,
        "{modelo}": modelo_formateado,
        "{inventario}": notebook.inventario or "Sin inventario",
        "{numero_serie}": notebook.numero_serie or "Sin número de serie",
        "{descripcion}": notebook.descripcion or "Sin descripción",
    }

    def reemplazar_texto(elemento):
        for placeholder, valor in placeholders.items():
            if placeholder in elemento.text:
                elemento.text = elemento.text.replace(placeholder, valor)

    for p in doc.paragraphs:
        reemplazar_texto(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_texto(cell)

    # Guardar el archivo generado
    archivo_path = os.path.join(UPLOAD_FOLDER, f"acta_entrega_{notebook.inventario}_{notebook.usuario}.docx")
    doc.save(archivo_path)

    return archivo_path

def generar_acta_devolucion(notebook):
    template_dir = os.path.abspath(UPLOAD_FOLDER)
    template_name = "acta_devolucion.docx"
    plantilla_path = os.path.join(template_dir, template_name)

    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"La plantilla de devolución '{plantilla_path}' no existe.")

    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

    doc = Document(plantilla_path)

    # Determinar el modelo formateado
    if notebook.modelo == 'dell i5':
        modelo_formateado = "Dell Latitude 3520 i5"
    elif notebook.modelo == 'dell i7':
        modelo_formateado = "Dell Latitude 5520 i7"
    elif notebook.modelo == 'dell i7 diseno':
        modelo_formateado = "Dell Precision 3561 i7"
    elif notebook.modelo == 'exo i5':
        modelo_formateado = "Exo Smart Pro Q6 i5"
    elif notebook.modelo == 'exo i7':
        modelo_formateado = "Exo Smart Pro Q6 i7"
    else:
        modelo_formateado = "Sin modelo"

    # Dirección formateada
    if notebook.direccion == 'intervencion':
        direccion_formateada = "Intervención"
    elif notebook.direccion == 'sistemas':
        direccion_formateada = "Dirección General de Sistemas Informáticos"
    elif notebook.direccion == 'fomento':
        direccion_formateada = "Dirección Nacional de Fomento y Desarrollo"
    elif notebook.direccion == 'administracion':
        direccion_formateada = "Dirección General de Administración"
    elif notebook.direccion == 'rrhh':
        direccion_formateada = "Dirección General de Recursos Humanos"
    elif notebook.direccion == 'juridicos':
        direccion_formateada = "Dirección General de Asuntos Jurídicos y Regulatorios"
    elif notebook.direccion == 'planificacion':
        direccion_formateada = "Dirección Nacional de Planificación y Convergencia"
    elif notebook.direccion == 'control':
        direccion_formateada = "Dirección Nacional de Control y Fiscalización"
    elif notebook.direccion == 'postales':
        direccion_formateada = "Dirección Nacional de Servicios Postales"
    elif notebook.direccion == 'institucionales':
        direccion_formateada = "Dirección General de Asuntos Institucionales"
    elif notebook.direccion == 'audiovisuales':
        direccion_formateada = "Dirección Nacional de Servicios Audiovisuales"
    elif notebook.direccion == 'competencia':
        direccion_formateada = "Dirección Nacional de Desarrollo de la Competencia en Redes"
    elif notebook.direccion == 'autorizaciones':
        direccion_formateada = "Dirección Nacional de Autorizaciones y Registros TIC"
    elif notebook.direccion == 'delegaciones':
        direccion_formateada = "Dirección Nacional de Atención de Usuarios"
    elif notebook.direccion == 'auditoria':
        direccion_formateada = "Unidad de Auditoría Interna"
    elif notebook.direccion == 'ccte':
        direccion_formateada = "Centro De Comprobación Técnica De Emisiones"
    else:
        direccion_formateada = "Sin dirección"


    # Diccionario de placeholders
    placeholders = {
        "{fecha}": notebook.fecha.strftime("%d/%m/%Y") if notebook.fecha else "Sin fecha",
        "{usuario}": notebook.usuario or "Sin usuario",
        "{direccion}": direccion_formateada,
        "{modelo}": modelo_formateado,
        "{inventario}": notebook.inventario or "Sin inventario",
        "{numero_serie}": notebook.numero_serie or "Sin número de serie",
        "{descripcion}": notebook.descripcion or "Sin descripción"
    }

    def reemplazar_texto(elemento):
        for placeholder, valor in placeholders.items():
            if placeholder in elemento.text:
                elemento.text = elemento.text.replace(placeholder, valor)

    for p in doc.paragraphs:
        reemplazar_texto(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_texto(cell)

    archivo_path = os.path.join(UPLOAD_FOLDER, f"acta_devolucion_{notebook.inventario}_{notebook.usuario}.docx")
    doc.save(archivo_path)

    return archivo_path

UPLOAD_FOLDER_PDF = os.path.join(os.getcwd(), 'uploads/actas_firmadas')
ALLOWED_EXTENSIONS = {'pdf'}

# Función para verificar la extensión del archivo
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@notebooks_bp.route('/uploads/actas_firmadas/<path:filename>')
@login_required
@lectura_allowed
def ver_pdf(filename):
    return send_from_directory('uploads/actas_firmadas', filename)

# Ruta para subir archivos
@notebooks_bp.route('/upload_pdf/<int:notebook_id>', methods=['POST'])
@login_required
@lectura_allowed
def upload_file(notebook_id):
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No se ha seleccionado ningún archivo.'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'success': False, 'message': 'No se ha seleccionado ningún archivo.'}), 400

    if file and allowed_file(file.filename):
        notebook = Notebook.query.get(notebook_id)
        if not notebook:
            return jsonify({'success': False, 'message': 'Notebook no encontrada.'}), 404

        # Usar directamente el nombre del usuario sin sanitizar
        filename = f"acta_firmada_{notebook.inventario}_{notebook.usuario.replace(' ', '_')}.pdf"
        ruta_completa = os.path.join(UPLOAD_FOLDER_PDF, filename)

        # Guardar archivo
        file.save(ruta_completa)

        # Guardar ruta relativa en la base de datos
        notebook.archivo_pdf = f"uploads/actas_firmadas/{filename}"
        db.session.commit()

        return jsonify({'success': True, 'message': 'Archivo subido correctamente.'}), 200

    return jsonify({'success': False, 'message': 'Solo se permiten archivos PDF.'}), 400


# Ruta para eliminar un archivo PDF
@notebooks_bp.route('/delete_pdf/<int:notebook_id>', methods=['POST'])
@login_required
@basico_required
def delete_pdf(notebook_id):
    notebook = Notebook.query.get_or_404(notebook_id)

    if notebook.archivo_pdf:
        # Construir ruta absoluta del archivo
        filename = f"acta_firmada_{notebook.inventario}_{notebook.usuario.replace(' ', '_')}.pdf"
        ruta_pdf = os.path.join(UPLOAD_FOLDER_PDF, filename)

        # Borrar el archivo físico si existe
        if os.path.exists(ruta_pdf):
            os.remove(ruta_pdf)

        # Borrar la referencia de la base de datos
        notebook.archivo_pdf = None
        db.session.commit()

        return jsonify({'success': True, 'message': 'PDF eliminado correctamente.'})
    
    return jsonify({'success': False, 'message': 'No hay archivo PDF asociado.'}), 404

# Ruta para servir archivos PDF con soporte para solicitudes de rango
@notebooks_bp.route('/uploads/actas_firmadas/<filename>')
@login_required
@lectura_allowed
def uploaded_file(filename):
    file_path = os.path.join(UPLOAD_FOLDER_PDF, filename)
    
    # Verificar si el archivo existe
    if not os.path.exists(file_path):
        return "Archivo no encontrado", 404
    
    # Manejar solicitudes de rango
    range_header = request.headers.get('Range')
    if not range_header:
        return send_from_directory(UPLOAD_FOLDER_PDF, filename)
    
    file_size = os.path.getsize(file_path)
    start, end = range_header.replace('bytes=', '').split('-')
    start = int(start)
    end = int(end) if end else file_size - 1
    
    with open(file_path, 'rb') as f:
        f.seek(start)
        chunk = f.read(end - start + 1)
    
    response = Response(chunk, 206, mimetype='application/pdf', direct_passthrough=True)
    response.headers.add('Content-Range', f'bytes {start}-{end}/{file_size}')
    response.headers.add('Accept-Ranges', 'bytes')
    response.headers.add('Content-Length', str(end - start + 1))
    
    return response