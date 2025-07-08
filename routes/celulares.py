from flask import Blueprint, request, render_template, redirect, current_app, send_file, send_from_directory, Response, jsonify
from flask_login import login_required, current_user
import unicodedata
from sqlalchemy import func
from models import db, Celular
from .auth import lectura_allowed, basico_required
from .logs import logger
from docx import Document
import os
from datetime import date

celulares_bp = Blueprint('celulares', __name__)

UPLOAD_FOLDER = os.path.abspath('uploads/actas_celulares')

def quitar_acentos(cadena):
    return ''.join(
        c for c in unicodedata.normalize('NFD', cadena)
        if unicodedata.category(c) != 'Mn'
    )

def safe_log(action, entity, details, user=None):
    """Función de logging ultra-resistente a fallos"""
    try:
        user = user or (current_user.username if current_user.is_authenticated else 'system')
        details_str = str(details)

        current_app.logger.info(f"Intentando registrar log: {action}, {entity}, {details_str}, {user}")

        if hasattr(logger, 'is_initialized') and logger.is_initialized:
            logger.log(action, entity, details_str, user)
        else:
            current_app.logger.warning("Logger personalizado no inicializado")
    except Exception as e:
        current_app.logger.error(f"Fallo en safe_log: {str(e)}")

@celulares_bp.route('/celulares', methods=['GET'])
@login_required
@lectura_allowed
def celulares():
    filtro = request.args.get('filtro', '')
    valor = request.args.get('valor', '')
    filtro_estado_celular = request.args.get('estado', '')
    filtro_modelo = request.args.get('modelo', '')
    filtro_direccion = request.args.get('direccion', '')

    # Construimos la consulta base
    consulta = Celular.query

    if filtro_estado_celular:
        consulta = consulta.filter(Celular.estado == filtro_estado_celular)

    if filtro_modelo:
        consulta = consulta.filter(Celular.modelo == filtro_modelo)

    if filtro_direccion:
        consulta = consulta.filter(Celular.direccion == filtro_direccion)
        
    if filtro and valor:
        valor_normalizado = quitar_acentos(valor).lower()

        if filtro == 'inventario':
            consulta = consulta.filter(Celular.inventario.contains(valor))
        elif filtro == 'imei':
            consulta = consulta.filter(Celular.imei.contains(valor))
        elif filtro == 'descripcion':
            consulta = consulta.filter(Celular.descripcion.contains(valor))
        elif filtro == 'usuario':
            valor_normalizado = quitar_acentos(valor).lower()
            consulta = consulta.filter(
                func.lower(func.replace(func.replace(func.replace(
                    func.replace(func.replace(func.replace(func.replace(
                        Celular.usuario,
                        'á', 'a'),
                        'é', 'e'),
                        'í', 'i'),
                        'ó', 'o'),
                        'ú', 'u'),
                        'Á', 'a'),
                        'É', 'e')
                ).contains(valor_normalizado)
            )

    celulares = consulta.all()
    return render_template('celulares.html', celulares=celulares, filtro=filtro, valor=valor, filtro_estado_celular=filtro_estado_celular, filtro_modelo=filtro_modelo, filtro_direccion=filtro_direccion)

@celulares_bp.route('/agregar_celular', methods=['GET','POST'])
@login_required
@basico_required
def agregar_celular():
    if request.method == 'POST':
        try:
            form_data = {
                'modelo' : request.form.get('modelo', '').strip(),
                'inventario' : request.form.get('inventario', '').strip(),
                'imei' : request.form.get('imei', '').strip(),
                'estado' : request.form.get('estado', '').strip(),
                'usuario' : request.form.get('usuario', '').strip(),
                'direccion' : request.form.get('direccion', '').strip(),
                'descripcion' : request.form.get('descripcion', '').strip(),
                'fecha' : request.form.get('fecha', '').strip(),
            }

            fecha = date.fromisoformat(form_data['fecha']) if form_data['fecha'] else None
        
            celular_existente = Celular.query.filter_by(inventario=form_data['inventario']).first()
            if celular_existente:
                log_data = {
                    'inventario': form_data['inventario'],
                    'intento': form_data,
                    'mensaje': 'el celular ya existe en la base de datos'
                }
                safe_log('conflict', 'celular', log_data, current_user.username)
                return """
                <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
                <script>
                    document.addEventListener("DOMContentLoaded", function() {
                        Swal.fire({
                            title: 'Error',
                            text: 'El número de inventario ya existe.',
                            icon: 'error',
                            confirmButtonText: 'Aceptar'
                        }).then(() => {
                            window.opener.location.reload();  // Recargar la página principal
                            window.close();
                        });
                    });
                </script>
                """

            nuevo_celular = Celular(
                modelo=form_data['modelo'],
                inventario=form_data['inventario'], 
                imei=form_data['imei'], 
                estado=form_data['estado'], 
                usuario=form_data['usuario'], 
                fecha=fecha,
                direccion=form_data['direccion'],
                descripcion=form_data['descripcion'] or None
            )

            db.session.add(nuevo_celular)
            db.session.commit()

            archivo_path = llenar_plantilla_acta(nuevo_celular)
            nuevo_celular.archivo_celular = archivo_path

            safe_log(
                action='create',
                entity='celular',
                details={
                    'id': nuevo_celular.id,
                    'modelo' : nuevo_celular.modelo,
                    'inventario' : nuevo_celular.inventario,
                    'imei' : nuevo_celular.imei,
                    'estado' : nuevo_celular.estado,
                    'usuario' : nuevo_celular.usuario,
                    'fecha' : nuevo_celular.fecha,
                    'direccion' : nuevo_celular.direccion,
                    'descripcion' : nuevo_celular.descripcion
                },
                user=current_user.username
            )
            return """
            <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
            <script>
                document.addEventListener("DOMContentLoaded", function() {
                    Swal.fire({
                        title: 'Un Éxito',
                        text: 'celular agregada correctamente.',
                        icon: 'success',
                        confirmButtonText: 'Aceptar'
                    }).then(() => {
                        window.opener.location.reload();  // Recargar la página principal
                        window.close();  // Cerrar la ventana emergente
                    });
                });
            </script>
            """
        except Exception as e:
            db.session.rollback()
            
            current_app.logger.error(f"Error en agregar_celular: {str(e)}")
            safe_log(
                action='error',
                entity='celular',
                details={'error': str(e), 'operacion': 'agregar_celular'},
                user=current_user.username
            )

            return """
            <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
            <script>
                document.addEventListener("DOMContentLoaded", function() {
                    Swal.fire({
                        title: 'Error',
                        text: 'Ocurrió un error al agregar el celular.',
                        icon: 'error',
                        confirmButtonText: 'Aceptar'
                    }).then(() => {
                        window.opener.location.reload();
                        window.close();
                    });
                });
            </script>
            """
        
    return render_template('agregar_celular.html')

@celulares_bp.route('/eliminar_celular/<int:id>', methods=['POST'])
@login_required
@basico_required
def eliminar_celular(id):
    try:
        celular = Celular.query.get(id)

        safe_log(
            action='delete',
            entity='celular',
            details={
                'id': id,
                'modelo' : celular.modelo,
                'inventario' : celular.inventario,
                'imei': celular.imei,
                'estado' : celular.estado,
                'usuario' : celular.usuario,
                'fecha' : celular.fecha,
                'direccion' : celular.direccion,
                'descripcion' : celular.descripcion
            },
            user=current_user.username
        )

        db.session.delete(celular)
        db.session.commit()

        return redirect('/celulares')
    
    except Exception as e:
        db.session.rollback()
        safe_log(
            action='error',
            entity='celular',
            details={
                'operation': 'delete',
                'error': str(e),
                'celular_id': id
            },
            user=current_user.username
        )
        return "Error al eliminar el celular", 500

@celulares_bp.route('/editar_celular/<int:id>', methods=['GET'])
@login_required
@basico_required
def editar_celular(id):
    celular = Celular.query.get(id)
    return render_template('editar_celular.html', celular=celular)

@celulares_bp.route('/actualizar_celular/<int:id>', methods=['POST'])
@login_required
@basico_required
def actualizar_celular(id):
    celular = Celular.query.get(id)

    if request.method == 'POST':
        try:
            datos_antiguos = {
                'modelo' : celular.modelo,
                'inventario' : celular.inventario,
                'imei' : celular.imei,
                'estado' : celular.estado,
                'usuario' : celular.usuario,
                'fecha' : celular.fecha,
                'direccion' : celular.direccion,
                'descripcion' : celular.descripcion
            }

            celular.modelo = request.form['modelo']
            celular.inventario = request.form['inventario']
            celular.imei = request.form['imei']
            celular.estado = request.form['estado']
            celular.usuario = request.form['usuario']
            celular.direccion = request.form['direccion']
            celular.descripcion = request.form['descripcion']

            celular.fecha = request.form['fecha']

            if celular.fecha:
                celular.fecha = date.fromisoformat(celular.fecha)
            else:
                celular.fecha = None  # Si el campo está vacío, guarda NULL en la base de datos

            if celular.estado.lower() == "fisica":
                if celular.archivo_celular and os.path.exists(celular.archivo_celular):
                    try:
                        os.remove(celular.archivo_celular)
                    except Exception as e:
                        print(f"No se pudo eliminar el archivo anterior: {e}")

                # 1. Generar el acta con los datos actuales
                archivo_devolucion = generar_acta_devolucion(celular)

                # 2. Guardar la ruta del acta de devolución en la base de datos
                celular.archivo_celular = archivo_devolucion

                # 4. Vaciar los campos
                celular.usuario = ""
                celular.direccion = ""
                celular.descripcion = ""

                # 5. Guardar cambios en la base de datos
                db.session.commit()

                safe_log(
                    action='update',
                    entity='celular',
                    details={
                        'id': id,
                        'cambios': {
                            'antes': datos_antiguos,
                            'despues': {
                                'modelo' : celular.modelo,
                                'inventario' : celular.inventario,
                                'imei' : celular.imei,
                                'estado' : celular.estado,
                                'usuario' : celular.usuario,
                                'fecha' : celular.fecha,
                                'direccion' : celular.direccion,
                                'descripcion' : celular.descripcion
                            }
                        }
                    },
                    user=current_user.username
                )

                # 6. Retornar HTML con descarga automática y SweetAlert
                return f"""
                <html>
                <head>
                    <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
                </head>
                <body>
                    <script>
                        document.addEventListener("DOMContentLoaded", function() {{
                            Swal.fire({{
                                title: 'Un Éxito',
                                text: 'celular actualizado correctamente... Descargando acta de devolución...',
                                icon: 'success',
                                confirmButtonText: 'Aceptar'
                            }}).then(() => {{
                                // Iniciar descarga
                                const link = document.createElement('a');
                                link.href = '/ruta_de_descarga_directa/{celular.id}';
                                link.download = '';
                                document.body.appendChild(link);
                                link.click();
                                document.body.removeChild(link);

                                // Recargar y cerrar ventana
                                setTimeout(() => {{
                                    window.opener.location.reload();
                                    window.close();
                                }}, 5000);
                            }});
                        }});
                    </script>
                </body>
                </html>
                """

            else:
                if celular.archivo_celular and os.path.exists(celular.archivo_celular):
                    try:
                        os.remove(celular.archivo_celular)
                    except Exception as e:
                        print(f"No se pudo eliminar el archivo anterior: {e}")

                archivo_entrega = llenar_plantilla_acta(celular)
        
                # Guardar la nueva ruta en la base de datos
                celular.archivo_celular = archivo_entrega
                db.session.commit()

            safe_log(
                action='update',
                entity='celular',
                details={
                    'id': id,
                    'cambios': {
                        'antes': datos_antiguos,
                        'despues': {
                            'modelo' : celular.modelo,
                            'inventario' : celular.inventario,
                            'imei' : celular.imei,
                            'estado' : celular.estado,
                            'usuario' : celular.usuario,
                            'fecha' : celular.fecha,
                            'direccion' : celular.direccion,
                            'descripcion' : celular.descripcion
                        }
                    }
                },
                user=current_user.username
            )
            return """
            <script src="https://cdn.jsdelivr.net/npm/sweetalert2@11"></script>
            <script>
                document.addEventListener("DOMContentLoaded", function() {
                    Swal.fire({
                        title: 'Un Éxito',
                        text: 'celular actualizado correctamente.',
                        icon: 'success',
                        confirmButtonText: 'Aceptar'
                    }).then(() => {
                        window.opener.location.reload();  // Recargar la página principal
                        window.close();  // Cerrar la ventana emergente
                    });
                });
            </script>
            """

        except Exception as e:
            db.session.rollback()
            safe_log(
                action='error',
                entity='celular',
                details={
                    'operation': 'update',
                    'error': str(e),
                    'celular_id': id
                },
                user=current_user.username
            )
            return "Error al actualizar el celular", 500

    return redirect('/celulares')

@celulares_bp.route('/celulares/ver_acta_celular/<int:id>')
@login_required
@lectura_allowed
def ver_acta_celular(id):
    celular = Celular.query.get_or_404(id)
    return render_template('acta_celular.html', celular=celular)

@celulares_bp.route('/ruta_de_descarga_directa_celular/<int:id>')
@login_required
@lectura_allowed
def ruta_de_descarga_directa_celular(id):
    celular = Celular.query.get_or_404(id)

    if celular.archivo_celular and os.path.exists(celular.archivo_celular):
        return send_file(celular.archivo_celular, as_attachment=True)
    else:
        return "Archivo no encontrado", 404

@celulares_bp.route('/descargar_acta_celular/<int:id>')
@login_required
@lectura_allowed
def descargar_acta_celular(id):
    # Obtener el remito desde la base de datos
    celular = Celular.query.get_or_404(id)

    # Verificar si celular ya tiene un archivo generado
    if not celular.archivo_celular or not os.path.exists(celular.archivo_celular):
        archivo_path = llenar_plantilla_acta(celular)
        
        # Guardar la nueva ruta en la base de datos
        celular.archivo_celular = archivo_path
        db.session.commit()

    # Si el remito está marcado como "fisica", generar el remito de devolución
    if celular.estado.lower() == "fisica":
        archivo_devolucion = generar_acta_devolucion(celular)
        
        # Guardar la ruta del acta de devolución en la base de datos
        celular.archivo_celular = archivo_devolucion
        db.session.commit()

        return send_file(archivo_devolucion, as_attachment=True)

    return send_file(celular.archivo_celular, as_attachment=True)

def llenar_plantilla_acta(celular):
    # Ruta de la plantilla
    template_dir = os.path.abspath(UPLOAD_FOLDER)
    template_name = "acta_entrega.docx"
    plantilla_path = os.path.join(template_dir, template_name)

    # Verificar si la plantilla existe
    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"La plantilla '{plantilla_path}' no existe.")

    # Crear directorio de salida si no existe
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

    # Cargar la plantilla
    doc = Document(plantilla_path)

    # Determinar el modelo formateado
    if celular.modelo == 'samsung style':
        modelo_formateado = "Samsung Galaxy Style LTE 4G G3"
    elif celular.modelo == 'samsung a5':
        modelo_formateado = "Samsung Galaxy A5"
    elif celular.modelo == 'samsung S9':
        modelo_formateado = "Samsung S9"
    elif celular.modelo == 'moto g60s':
        modelo_formateado = "Moto G60S"
    elif celular.modelo == 'samsung note 8':
        modelo_formateado = "Samsung Note 8"
    elif celular.modelo == 'samsung a3':
        modelo_formateado = "Samsung Galaxy A3"
    elif celular.modelo == 'samsung grand prime':
        modelo_formateado = "Samsung Galaxy Grand Prime 4G"
    elif celular.modelo == 'samsung j2':
        modelo_formateado = "Samsung Galaxy J2 Prime"
    elif celular.modelo == 'moto x':
        modelo_formateado = "Moto X 2da Gen 4G LTE"
    elif celular.modelo == 'samsung a51':
        modelo_formateado = "Samsung A51"
    elif celular.modelo == 'samsung j1':
        modelo_formateado = "Samsung Galaxy J1 Ace"
    elif celular.modelo == 'samsung s7':
        modelo_formateado = "Samsung Galaxy S7 Edge"
    elif celular.modelo == 'moto x play':
        modelo_formateado = "Moto X Play"
    elif celular.modelo == 'samsung j5':
        modelo_formateado = "Samsung J5 Prime"
    else:
        modelo_formateado = "Sin modelo"

    # Dirección formateada
    if celular.direccion == 'intervencion':
        direccion_formateada = "Intervención"
    elif celular.direccion == 'sistemas':
        direccion_formateada = "Dirección General de Sistemas Informáticos"
    elif celular.direccion == 'fomento':
        direccion_formateada = "Dirección Nacional de Fomento y Desarrollo"
    elif celular.direccion == 'administracion':
        direccion_formateada = "Dirección General de Administración"
    elif celular.direccion == 'rrhh':
        direccion_formateada = "Dirección General de Recursos Humanos"
    elif celular.direccion == 'juridicos':
        direccion_formateada = "Dirección General de Asuntos Jurídicos y Regulatorios"
    elif celular.direccion == 'planificacion':
        direccion_formateada = "Dirección Nacional de Planificación y Convergencia"
    elif celular.direccion == 'control':
        direccion_formateada = "Dirección Nacional de Control y Fiscalización"
    elif celular.direccion == 'postales':
        direccion_formateada = "Dirección Nacional de Servicios Postales"
    elif celular.direccion == 'institucionales':
        direccion_formateada = "Dirección General de Asuntos Institucionales"
    elif celular.direccion == 'audiovisuales':
        direccion_formateada = "Dirección Nacional de Servicios Audiovisuales"
    elif celular.direccion == 'competencia':
        direccion_formateada = "Dirección Nacional de Desarrollo de la Competencia en Redes"
    elif celular.direccion == 'autorizaciones':
        direccion_formateada = "Dirección Nacional de Autorizaciones y Registros TIC"
    elif celular.direccion == 'delegaciones':
        direccion_formateada = "Dirección Nacional de Atención de Usuarios"
    elif celular.direccion == 'auditoria':
        direccion_formateada = "Unidad de Auditoría Interna"
    elif celular.direccion == 'ccte':
        direccion_formateada = "Centro De Comprobación Técnica De Emisiones"
    else:
        direccion_formateada = "Sin dirección"


    # Diccionario de placeholders
    placeholders = {
        "{fecha}": celular.fecha.strftime("%d/%m/%Y") if celular.fecha else "Sin fecha",
        "{usuario}": celular.usuario or "Sin usuario",
        "{direccion}": direccion_formateada,
        "{modelo}": modelo_formateado,
        "{inventario}": celular.inventario or "Sin inventario",
        "{imei}": celular.imei or "Sin imei",
        "{descripcion}": celular.descripcion or "Sin descripción",
    }

    def reemplazar_texto(elemento):
        for placeholder, valor in placeholders.items():
            if placeholder in elemento.text:
                elemento.text = elemento.text.replace(placeholder, valor)

    for p in doc.paragraphs:
        reemplazar_texto(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_texto(cell)

    # Guardar el archivo generado
    archivo_path = os.path.join(UPLOAD_FOLDER, f"acta_entrega_{celular.inventario}_{celular.usuario}.docx")
    doc.save(archivo_path)

    return archivo_path

def generar_acta_devolucion(celular):
    template_dir = os.path.abspath(UPLOAD_FOLDER)
    template_name = "acta_devolucion.docx"
    plantilla_path = os.path.join(template_dir, template_name)

    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"La plantilla de devolución '{plantilla_path}' no existe.")

    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

    doc = Document(plantilla_path)

    # Determinar el modelo formateado
    if celular.modelo == 'samsung style':
        modelo_formateado = "Samsung Galaxy Style LTE 4G G3"
    elif celular.modelo == 'samsung a5':
        modelo_formateado = "Samsung Galaxy A5"
    elif celular.modelo == 'samsung S9':
        modelo_formateado = "Samsung S9"
    elif celular.modelo == 'moto g60s':
        modelo_formateado = "Moto G60S"
    elif celular.modelo == 'samsung note 8':
        modelo_formateado = "Samsung Note 8"
    elif celular.modelo == 'samsung a3':
        modelo_formateado = "Samsung Galaxy A3"
    elif celular.modelo == 'samsung grand prime':
        modelo_formateado = "Samsung Galaxy Grand Prime 4G"
    elif celular.modelo == 'samsung j2':
        modelo_formateado = "Samsung Galaxy J2 Prime"
    elif celular.modelo == 'moto x':
        modelo_formateado = "Moto X 2da Gen 4G LTE"
    elif celular.modelo == 'samsung a51':
        modelo_formateado = "Samsung A51"
    elif celular.modelo == 'samsung j1':
        modelo_formateado = "Samsung Galaxy J1 Ace"
    elif celular.modelo == 'samsung s7':
        modelo_formateado = "Samsung Galaxy S7 Edge"
    elif celular.modelo == 'moto x play':
        modelo_formateado = "Moto X Play"
    elif celular.modelo == 'samsung j5':
        modelo_formateado = "Samsung J5 Prime"
    else:
        modelo_formateado = "Sin modelo"

    # Dirección formateada
    if celular.direccion == 'intervencion':
        direccion_formateada = "Intervención"
    elif celular.direccion == 'sistemas':
        direccion_formateada = "Dirección General de Sistemas Informáticos"
    elif celular.direccion == 'fomento':
        direccion_formateada = "Dirección Nacional de Fomento y Desarrollo"
    elif celular.direccion == 'administracion':
        direccion_formateada = "Dirección General de Administración"
    elif celular.direccion == 'rrhh':
        direccion_formateada = "Dirección General de Recursos Humanos"
    elif celular.direccion == 'juridicos':
        direccion_formateada = "Dirección General de Asuntos Jurídicos y Regulatorios"
    elif celular.direccion == 'planificacion':
        direccion_formateada = "Dirección Nacional de Planificación y Convergencia"
    elif celular.direccion == 'control':
        direccion_formateada = "Dirección Nacional de Control y Fiscalización"
    elif celular.direccion == 'postales':
        direccion_formateada = "Dirección Nacional de Servicios Postales"
    elif celular.direccion == 'institucionales':
        direccion_formateada = "Dirección General de Asuntos Institucionales"
    elif celular.direccion == 'audiovisuales':
        direccion_formateada = "Dirección Nacional de Servicios Audiovisuales"
    elif celular.direccion == 'competencia':
        direccion_formateada = "Dirección Nacional de Desarrollo de la Competencia en Redes"
    elif celular.direccion == 'autorizaciones':
        direccion_formateada = "Dirección Nacional de Autorizaciones y Registros TIC"
    elif celular.direccion == 'delegaciones':
        direccion_formateada = "Dirección Nacional de Atención de Usuarios"
    elif celular.direccion == 'auditoria':
        direccion_formateada = "Unidad de Auditoría Interna"
    elif celular.direccion == 'ccte':
        direccion_formateada = "Centro De Comprobación Técnica De Emisiones"
    else:
        direccion_formateada = "Sin dirección"


    # Diccionario de placeholders
    placeholders = {
        "{fecha}": celular.fecha.strftime("%d/%m/%Y") if celular.fecha else "Sin fecha",
        "{usuario}": celular.usuario or "Sin usuario",
        "{direccion}": direccion_formateada,
        "{modelo}": modelo_formateado,
        "{inventario}": celular.inventario or "Sin inventario",
        "{imei}": celular.imei or "Sin imei",
        "{descripcion}": celular.descripcion or "Sin descripción",
    }

    def reemplazar_texto(elemento):
        for placeholder, valor in placeholders.items():
            if placeholder in elemento.text:
                elemento.text = elemento.text.replace(placeholder, valor)

    for p in doc.paragraphs:
        reemplazar_texto(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_texto(cell)

    archivo_path = os.path.join(UPLOAD_FOLDER, f"acta_devolucion_{celular.inventario}_{celular.usuario}.docx")
    doc.save(archivo_path)

    return archivo_path

UPLOAD_FOLDER_PDF = os.path.join(os.getcwd(), 'uploads/actas_firmadas_cel')
ALLOWED_EXTENSIONS = {'pdf'}

# Función para verificar la extensión del archivo
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@celulares_bp.route('/uploads/actas_firmadas_cel/<path:filename>')
@login_required
@lectura_allowed
def ver_pdf_celular(filename):
    return send_from_directory('uploads/actas_firmadas_cel', filename)

# Ruta para subir archivos
@celulares_bp.route('/upload_pdf_celular/<int:celular_id>', methods=['POST'])
@login_required
@lectura_allowed
def upload_file_celular(celular_id):
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No se ha seleccionado ningún archivo.'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'success': False, 'message': 'No se ha seleccionado ningún archivo.'}), 400

    if file and allowed_file(file.filename):
        celular = Celular.query.get(celular_id)
        if not celular:
            return jsonify({'success': False, 'message': 'celular no encontrada.'}), 404

        # Usar directamente el nombre del usuario sin sanitizar
        filename = f"acta_firmada_{celular.inventario}_{celular.usuario.replace(' ', '_')}.pdf"
        ruta_completa = os.path.join(UPLOAD_FOLDER_PDF, filename)

        # Guardar archivo
        file.save(ruta_completa)

        # Guardar ruta relativa en la base de datos
        celular.archivo_pdf_c = f"uploads/actas_firmadas_cel/{filename}"
        db.session.commit()

        return jsonify({'success': True, 'message': 'Archivo subido correctamente.'}), 200

    return jsonify({'success': False, 'message': 'Solo se permiten archivos PDF.'}), 400


# Ruta para eliminar un archivo PDF
@celulares_bp.route('/delete_pdf_celular/<int:celular_id>', methods=['POST'])
@login_required
@basico_required
def delete_pdf_celular(celular_id):
    celular = Celular.query.get_or_404(celular_id)

    if celular.archivo_pdf_c:
        # Construir ruta absoluta del archivo
        filename = f"acta_firmada_{celular.inventario}_{celular.usuario.replace(' ', '_')}.pdf"
        ruta_pdf = os.path.join(UPLOAD_FOLDER_PDF, filename)

        # Borrar el archivo físico si existe
        if os.path.exists(ruta_pdf):
            os.remove(ruta_pdf)

        # Borrar la referencia de la base de datos
        celular.archivo_pdf_c = None
        db.session.commit()

        return jsonify({'success': True, 'message': 'PDF eliminado correctamente.'})
    
    return jsonify({'success': False, 'message': 'No hay archivo PDF asociado.'}), 404

# Ruta para servir archivos PDF con soporte para solicitudes de rango
@celulares_bp.route('/uploads/actas_firmadas_cel/<filename>')
@login_required
@lectura_allowed
def uploaded_file_celular(filename):
    file_path = os.path.join(UPLOAD_FOLDER_PDF, filename)
    
    # Verificar si el archivo existe
    if not os.path.exists(file_path):
        return "Archivo no encontrado", 404
    
    # Manejar solicitudes de rango
    range_header = request.headers.get('Range')
    if not range_header:
        return send_from_directory(UPLOAD_FOLDER_PDF, filename)
    
    file_size = os.path.getsize(file_path)
    start, end = range_header.replace('bytes=', '').split('-')
    start = int(start)
    end = int(end) if end else file_size - 1
    
    with open(file_path, 'rb') as f:
        f.seek(start)
        chunk = f.read(end - start + 1)
    
    response = Response(chunk, 206, mimetype='application/pdf', direct_passthrough=True)
    response.headers.add('Content-Range', f'bytes {start}-{end}/{file_size}')
    response.headers.add('Accept-Ranges', 'bytes')
    response.headers.add('Content-Length', str(end - start + 1))
    
    return response